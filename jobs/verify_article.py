#!/usr/bin/env python3
"""Article verification CLI — проверяет статью и базу знаний на ошибки.

Usage:
    python3 -m jobs.verify_article
    python3 -m jobs.verify_article --docx /path/to/RU_FINAL_v7.docx
    python3 -m jobs.verify_article --rules-only
    python3 -m jobs.verify_article --print-rules

Что проверяется:
  1. База знаний (production_rules_knowledge_base.json):
     - Уникальность подписей (sorted(antecedent), consequent, condition).
     - Все индексы признаков ∈ SHAP top-15.
     - Консистентность антецедентов с условиями.
     - Распределение классов (a₁ / a₂ / a₃) и размеров антецедентов.
     - Использование каждого top-15 признака.
  2. Финальный DOCX:
     - Подсчёт реальных правил-параграфов p1...p150.
     - Поиск orphan OMML-формул (старые "p100" и т.п.).
     - Поиск остаточных "100 правил" в тексте и таблицах.
     - Подписи рисунков 1–9 (последовательность, отсутствие дублей).
     - Все ссылки "Формула (X)" указывают на существующие формулы.
     - Длина аннотации (≤200 слов для JDMSC).
     - Наличие AMS Subject Classification.
     - Ключевые формулы (12) и (17) корректны.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = Path("/Users/asadbekikromov/Downloads/Магистратура/mine/Q1/2/output_final/RU_FINAL_v7.docx")
KB_JSON = REPO / "production_rules_knowledge_base.json"


# ─────────────────────── ANSI colours for CLI output ───────────────────────
class C:
    OK = "\033[92m"
    WARN = "\033[93m"
    ERR = "\033[91m"
    BOLD = "\033[1m"
    DIM = "\033[2m"
    HEAD = "\033[96m"
    END = "\033[0m"


def hdr(msg: str) -> None:
    print(f"\n{C.HEAD}{C.BOLD}{'=' * 70}{C.END}")
    print(f"{C.HEAD}{C.BOLD}  {msg}{C.END}")
    print(f"{C.HEAD}{C.BOLD}{'=' * 70}{C.END}")


def ok(msg: str) -> None:
    print(f"  {C.OK}✓{C.END} {msg}")


def warn(msg: str) -> None:
    print(f"  {C.WARN}⚠{C.END} {msg}")


def err(msg: str) -> None:
    print(f"  {C.ERR}✗{C.END} {msg}")


# ─────────────────────── KB verification ───────────────────────
def check_kb(print_rules: bool = False) -> tuple[int, int]:
    """Return (warnings, errors)."""
    hdr("1. ПРОВЕРКА БАЗЫ ЗНАНИЙ (production_rules_knowledge_base.json)")

    if not KB_JSON.exists():
        err(f"Файл не найден: {KB_JSON}")
        return 0, 1

    kb = json.loads(KB_JSON.read_text(encoding="utf-8"))
    rules = kb.get("rules", [])
    top15 = set(kb.get("shap_top15_indices", []))
    events = set(kb.get("events", {}).keys())

    print(f"{C.DIM}  Версия: v{kb.get('version','?')} • Всего правил: {len(rules)} • SHAP top-15 индексов: {len(top15)}{C.END}")

    warnings = errors = 0

    # 1.1 Uniqueness
    sigs = [(tuple(sorted(r["ante"])), r["cons"], r["cond"]) for r in rules]
    dup_count = len(sigs) - len(set(sigs))
    if dup_count == 0:
        ok(f"Уникальность: {len(set(sigs))} / {len(sigs)} — нет дублей")
    else:
        err(f"Уникальность: {dup_count} дубликатов!")
        errors += 1

    # 1.2 All antecedents in top-15
    bad = []
    for i, r in enumerate(rules, 1):
        for a in r["ante"]:
            if a not in top15:
                bad.append((i, a))
    if not bad:
        ok(f"Все антецеденты в SHAP top-15 ({sorted(top15)})")
    else:
        err(f"Антецеденты вне top-15: {len(bad)}; пример: {bad[:3]}")
        errors += 1

    # 1.3 Valid consequents
    bad_cons = [(i + 1, r["cons"]) for i, r in enumerate(rules) if r["cons"] not in events]
    if not bad_cons:
        ok(f"Все консеквенты валидны ({events})")
    else:
        err(f"Невалидных консеквентов: {len(bad_cons)}")
        errors += 1

    # 1.4 Antecedent <-> condition match
    mism = 0
    for i, r in enumerate(rules, 1):
        cond_idx = set(int(m) for m in re.findall(r"f(\d+)", r["cond"]))
        if cond_idx != set(r["ante"]):
            mism += 1
    if mism == 0:
        ok("Антецеденты соответствуют условиям (нет рассогласований)")
    else:
        err(f"Рассогласований ante↔condition: {mism}")
        errors += 1

    # 1.5 Class balance
    cnts = Counter(r["cons"] for r in rules)
    print(f"{C.DIM}  Распределение классов: a1={cnts.get('a1',0)}, a2={cnts.get('a2',0)}, a3={cnts.get('a3',0)}{C.END}")

    # 1.6 Antecedent size
    sizes = Counter(len(r["ante"]) for r in rules)
    print(f"{C.DIM}  Размеры антецедентов: 3={sizes.get(3,0)}, 4={sizes.get(4,0)}, 5={sizes.get(5,0)}{C.END}")

    # 1.7 Feature usage
    usage = Counter()
    for r in rules:
        for a in r["ante"]:
            usage[a] += 1
    if set(usage) >= top15:
        min_u = min(usage[i] for i in top15)
        max_u = max(usage[i] for i in top15)
        ok(f"Все 15 признаков задействованы (min={min_u}, max={max_u}, mean={sum(usage.values())/len(top15):.1f})")
    else:
        unused = top15 - set(usage)
        warn(f"Неиспользуемые признаки: {sorted(unused)}")
        warnings += 1

    if print_rules:
        print(f"\n{C.BOLD}  ВСЕ ПРАВИЛА:{C.END}")
        sub_digits = "₀₁₂₃₄₅₆₇₈₉"

        def f_label(idx):
            return "f" + "".join(sub_digits[int(c)] for c in str(idx))

        for i, r in enumerate(rules, 1):
            ante = " ∧ ".join(f_label(j) for j in r["ante"])
            cond = re.sub(r"f(\d+)", lambda m: f_label(int(m.group(1))), r["cond"])
            print(f"    p{i:3d}: {ante} → {r['cons']} | {cond}")

    return warnings, errors


# ─────────────────────── DOCX verification ───────────────────────
def check_docx(docx_path: Path) -> tuple[int, int]:
    """Return (warnings, errors)."""
    hdr(f"2. ПРОВЕРКА DOCX ({docx_path.name})")

    if not docx_path.exists():
        err(f"Файл не найден: {docx_path}")
        return 0, 1

    try:
        from docx import Document
    except ImportError:
        err("Не установлен python-docx. Установи: pip3 install python-docx")
        return 0, 1

    d = Document(docx_path)
    M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    print(f"{C.DIM}  Параграфов: {len(d.paragraphs)} • Таблиц: {len(d.tables)} • Размер: {docx_path.stat().st_size:,} байт{C.END}")

    warnings = errors = 0

    # 2.1 Real rule paragraphs
    rule_paras = []
    for i, p in enumerate(d.paragraphs):
        m = re.match(r"^p(\d+)\s*:", p.text.strip())
        if m:
            rule_paras.append((i, int(m.group(1))))
    nums = sorted(r[1] for r in rule_paras)
    expected_max = nums[-1] if nums else 0
    gaps = [n for n in range(1, expected_max + 1) if n not in nums]
    dupes = [n for n in set(nums) if nums.count(n) > 1]
    if rule_paras and not gaps and not dupes:
        ok(f"{len(rule_paras)} правил-параграфов p1..p{expected_max} (без пропусков и дублей)")
    else:
        err(f"Правил: {len(rule_paras)}, пропуски: {gaps[:5]}, дубликаты: {dupes[:5]}")
        errors += 1

    # 2.2 Orphan rule-OMML
    n_orphan = 0
    for i, p in enumerate(d.paragraphs):
        plain = p.text.strip()
        is_real_rule = bool(plain) and re.match(r"^p\d+\s*:", plain)
        if is_real_rule:
            continue
        for om in p._element.findall(f".//{{{M_NS}}}oMath"):
            text = "".join(t.text or "" for t in om.iter(f"{{{M_NS}}}t"))
            if re.search(r"p[\d_{}]+\s*:.*→\s*a", text):
                n_orphan += 1
                print(f"    {C.ERR}↳ orphan at p{i}: {text[:100]}{C.END}")
    if n_orphan == 0:
        ok("Orphan OMML-правил: 0")
    else:
        err(f"Orphan OMML-правил: {n_orphan}")
        errors += 1

    # 2.3 Old "100 правил" remnants
    old100 = []
    patterns = ["100 правил", "100 продукционных правил", "100 уникальных правил"]
    for i, p in enumerate(d.paragraphs):
        for pat in patterns:
            if pat in p.text:
                old100.append((f"p{i}", p.text[:120]))
    for ti, t in enumerate(d.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pat in patterns:
                    if pat in cell.text:
                        old100.append((f"T{ti+1}r{ri}c{ci}", cell.text[:120]))
    if not old100:
        ok("Остатков '100 правил' / '100 продукционных правил': 0")
    else:
        err(f"Остатков '100 правил': {len(old100)}")
        for loc, txt in old100[:3]:
            print(f"    {C.ERR}↳ [{loc}] {txt}{C.END}")
        errors += 1

    # 2.4 Figure captions sequence
    captions = []
    for p in d.paragraphs:
        m = re.match(r"^Рис\.\s*(\d+)\.", p.text.strip())
        if m:
            captions.append(int(m.group(1)))
    img_count = sum(1 for p in d.paragraphs if p._element.findall(f".//{{{W_NS}}}drawing"))
    cap_dupes = [n for n in set(captions) if captions.count(n) > 1]
    if captions == list(range(1, 10)) and img_count == 9:
        ok("Подписи Рис. 1–9 последовательны, 9 встроенных рисунков")
    else:
        err(f"Подписи: {captions}, дубликаты: {cap_dupes}, картинок: {img_count}")
        errors += 1

    # 2.5 Formula references
    text_refs = set()
    for p in d.paragraphs:
        for m in re.finditer(r"Формул[аы]\s*\((\d+)\)", p.text):
            text_refs.add(int(m.group(1)))
        for m in re.finditer(r"Формулы\s*(\d+)[–-](\d+)", p.text):
            for n in range(int(m.group(1)), int(m.group(2)) + 1):
                text_refs.add(n)
    omml_nums = set()
    for p in d.paragraphs:
        for om in p._element.findall(f".//{{{M_NS}}}oMath"):
            text = "".join(t.text or "" for t in om.iter(f"{{{M_NS}}}t"))
            for m in re.finditer(r"\((\d+)\)", text):
                omml_nums.add(int(m.group(1)))
    missing = text_refs - omml_nums
    if not missing:
        ok(f"Все ссылки 'Формула (X)' валидны (есть {len(omml_nums)} формул)")
    else:
        err(f"Несуществующих формул в ссылках: {sorted(missing)}")
        errors += 1

    # 2.6 Abstract length
    abs_p = next((p for p in d.paragraphs if p.text.strip().startswith("Аннотация")), None)
    if abs_p:
        words = len(abs_p.text.split())
        if words <= 200:
            ok(f"Аннотация: {words} слов (≤200 для JDMSC)")
        else:
            warn(f"Аннотация {words} слов > 200 — длинновата для JDMSC")
            warnings += 1

    # 2.7 AMS classification
    ams = next((p for p in d.paragraphs if "AMS" in p.text and "Subject" in p.text), None)
    if ams:
        ok("AMS Mathematics Subject Classification присутствует")
    else:
        warn("AMS Subject Classification не найдена")
        warnings += 1

    # 2.8 Key formulas
    f12_ok = False
    f17_ok = False
    for p in d.paragraphs[:90]:
        for om in p._element.findall(f".//{{{M_NS}}}oMath"):
            text = "".join(t.text or "" for t in om.iter(f"{{{M_NS}}}t"))
            if "(12)" in text and "|F| = 56" in text:
                f12_ok = True
            if "(17)" in text and "1·p̂_{LR}" in text and "1·p̂_{RF}" in text and "/ 9" in text:
                f17_ok = True
    if f12_ok:
        ok("Формула (12): |F| = 56 — корректна")
    else:
        err("Формула (12) не содержит '|F| = 56'")
        errors += 1
    if f17_ok:
        ok("Формула (17): 5-моделей HYBRID_VOTING / 9 — корректна")
    else:
        err("Формула (17) не содержит 5-модельной записи / 9")
        errors += 1

    return warnings, errors


# ─────────────────────── main ───────────────────────
def main():
    ap = argparse.ArgumentParser(description="Verify Q1 article + production rules KB.")
    ap.add_argument("--docx", default=str(DEFAULT_DOCX), help="Path to final DOCX")
    ap.add_argument("--rules-only", action="store_true", help="Only check the rules KB")
    ap.add_argument("--docx-only", action="store_true", help="Only check the DOCX")
    ap.add_argument("--print-rules", action="store_true", help="Print all 150 rules to stdout")
    args = ap.parse_args()

    total_w = total_e = 0
    if not args.docx_only:
        w, e = check_kb(print_rules=args.print_rules)
        total_w += w
        total_e += e
    if not args.rules_only:
        w, e = check_docx(Path(args.docx))
        total_w += w
        total_e += e

    print()
    hdr("ИТОГ")
    print(f"  {C.WARN if total_w else C.OK}Предупреждения:{C.END} {total_w}")
    print(f"  {C.ERR if total_e else C.OK}Ошибки:{C.END} {total_e}")
    if total_e == 0:
        print(f"\n  {C.OK}{C.BOLD}✅ ВСЁ ОК — статья и база знаний валидны.{C.END}")
        return 0
    else:
        print(f"\n  {C.ERR}{C.BOLD}❌ Найдены ошибки — нужно исправить.{C.END}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
